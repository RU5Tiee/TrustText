import os
import re
import html
import unicodedata
import uuid
import datetime
from collections import Counter
from bs4 import BeautifulSoup
import spacy

from schema import PolicyDocument, PolicyMetadata, Section, Clause, Entity

# =============================================================================
# CONSTANTS
# =============================================================================

# Known abbreviations that end with '.' but must NOT trigger sentence splits
ABBREVS = {
    'u.s', 'u.k', 'u.s.a', 'e.g', 'i.e', 'etc', 'vs', 'cf',
    'dr', 'mr', 'mrs', 'ms', 'prof', 'sr', 'jr', 'rev',
    'inc', 'corp', 'ltd', 'llc', 'co', 'plc', 'llp',
    'no', 'nos', 'sec', 'art', 'para', 'fig', 'dept', 'div',
    'govt', 'approx', 'est', 'ref', 'vol', 'pp', 'op',
    'jan', 'feb', 'mar', 'apr', 'jun', 'jul', 'aug',
    'sep', 'oct', 'nov', 'dec',
}

# HTML tags that are definitively noise — never contain policy text
NOISE_TAGS = [
    "script", "style", "nav", "footer", "header", "aside", "noscript",
    "form", "button", "figure", "figcaption", "iframe", "svg", "canvas",
    "video", "audio", "picture", "template", "dialog", "menu", "menuitem",
]

# CSS class / id fragments that identify noise containers
NOISE_CLASS_RE = re.compile(
    r'cookie|consent|gdpr|banner|popup|modal|overlay|tooltip|dropdown|'
    r'advert|advertisement|\bad\b|ad-unit|sidebar|widget|breadcrumb|'
    r'pagination|share|social|tweet|facebook|print|subscribe|newsletter|'
    r'navigation|navbar|topbar|masthead|eyebrow|ribbon|flyout|drawer|'
    r'toast|notification|alert-bar|sticky|floating|back-to-top|'
    r'skip-to|skip-link|promo|callout|cta|related|recommended|'
    r'trending|most-popular|load-more|show-more|see-all',
    re.I,
)

# Regex patterns to strip noise BEFORE segmentation
_NOISE_PATTERNS = [
    # Timestamps / effective-date lines
    re.compile(
        r'(last\s+updated?|effective\s+date|revised?\s+on?|updated?\s+on|'
        r'version|published)\s*:?\s*[^\n]{0,80}\d{4}',
        re.I,
    ),
    # Copyright lines
    re.compile(r'©\s*\d{4}[^\n]*', re.I),
    re.compile(r'copyright\s*©?\s*\d{4}[^\n]*', re.I),
    re.compile(r'all\s+rights\s+reserved[^\n]*', re.I),
    # Date formats
    re.compile(r'\b\d{1,2}[/\-\.]\d{1,2}[/\-\.]\d{2,4}\b'),
    re.compile(r'\b\d{4}[/\-\.]\d{2}[/\-\.]\d{2}\b'),
    re.compile(
        r'\b(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*'
        r'\.?\s+\d{1,2},?\s+\d{4}\b',
        re.I,
    ),
    # Bare URLs and emails
    re.compile(r'https?://\S+'),
    re.compile(r'www\.\S+\.\S+'),
    re.compile(r'[\w\.\+\-]+@[\w\.\-]+\.\w{2,}'),
    # Nav breadcrumbs:  Home > Legal > Privacy
    re.compile(r'(\w[\w\s]{1,30})\s*[>›|]\s*(\w[\w\s]{1,30})'),
    # Isolated nav link lines (only common words, no sentence verbs)
    re.compile(
        r'^\s*(home|about|contact|terms|privacy|sitemap|faq|help|'
        r'careers|press|blog|login|sign\s*up|register|account)'
        r'(\s*\|?\s*(home|about|contact|terms|privacy|sitemap|faq|'
        r'help|careers|press|blog|login|sign\s*up|register|account))+\s*$',
        re.I | re.M,
    ),
    # Action buttons bled in
    re.compile(
        r'^\s*(share|tweet|pin|print|email|download|back to top|'
        r'read more|show more|load more|close|dismiss|accept|decline)\s*$',
        re.I | re.M,
    ),
    # Page numbers
    re.compile(r'^\s*page\s+\d+\s*(of\s*\d+)?\s*$', re.I | re.M),
    re.compile(r'^\s*\d+\s*$', re.M),
]

# Sentence boundary: split on ". " followed by an uppercase letter,
# but only when the character before '.' is a lowercase letter, digit, or closing punct.
_SENT_SPLIT_RE = re.compile(r'(?<=[a-z0-9\"\'\)])\.\s+(?=[A-Z\"\(])')

# Heading heuristic: short (≤10 words), starts uppercase, no trailing period,
# optionally numbered  "3. Data Sharing"
_HEADING_RE = re.compile(r'^(\d+[\.\)]\s+)?[A-Z\u201C][^\n]{0,79}$')


# =============================================================================
# HELPERS
# =============================================================================

def _split_sentences(text: str) -> list:
    """
    Split text into individual sentences using the fullstop heuristic.
    Abbreviation guard: if the word immediately before the split is a known
    abbreviation, the two fragments are stitched back together.
    """
    raw_parts = _SENT_SPLIT_RE.split(text)
    sentences = []
    for part in raw_parts:
        part = part.strip()
        if not part:
            continue
        tokens = re.split(r'\s+', part)
        last_word = tokens[-1].rstrip('.').lower() if tokens else ''
        if last_word in ABBREVS and sentences:
            sentences[-1] = sentences[-1].rstrip() + '. ' + part
        else:
            sentences.append(part)
    return sentences


def _group_sentences(sentences: list, window: int = 3, max_words: int = 200) -> list:
    """
    Group consecutive sentences into clause-sized chunks of ≤ max_words.
    If a group exceeds max_words, reduce window by 1 and retry.
    """
    clauses = []
    i = 0
    while i < len(sentences):
        w = window
        while w > 0:
            chunk = sentences[i: i + w]
            combined = ' '.join(chunk).strip()
            if len(combined.split()) <= max_words or w == 1:
                clauses.append(combined)
                i += w
                break
            w -= 1
    return clauses


# =============================================================================
# 1. TEXT CLEANER
# =============================================================================

class TextCleaner:
    @staticmethod
    def clean(text: str) -> str:
        # Decode HTML entities and normalise unicode
        text = html.unescape(text)
        text = unicodedata.normalize('NFKD', text)

        # Strip all known noise patterns
        for pat in _NOISE_PATTERNS:
            text = pat.sub(' ', text)

        # Drop lines that carry fewer than 3 real alphabetic words
        # (catches stray nav links, lone punctuation, icon labels, etc.)
        clean_lines = []
        for line in text.splitlines():
            stripped = line.strip()
            alpha_words = re.findall(r'\b[a-zA-Z]{2,}\b', stripped)
            if len(alpha_words) >= 3:
                clean_lines.append(line)
            else:
                clean_lines.append('')   # preserve paragraph-break potential
        text = '\n'.join(clean_lines)

        # Keep only printable characters + essential punctuation
        text = re.sub(r'[^\w\s\.,;:\-\'\"\(\)\[\]\{\}?!@#\$%&\*\/\n]', ' ', text)

        # Collapse horizontal whitespace, then excessive blank lines
        text = re.sub(r'[ \t]+', ' ', text)
        text = re.sub(r'\n{3,}', '\n\n', text)

        # Deduplicate repeated lines (footer / cookie boilerplate echoed on every page)
        seen: set = set()
        deduped = []
        for line in text.splitlines():
            key = line.strip().lower()
            if key and key in seen:
                continue
            seen.add(key)
            deduped.append(line)
        text = '\n'.join(deduped)

        return text.strip()


# =============================================================================
# 2. ADAPTERS
# =============================================================================

class BaseAdapter:
    def parse(self, filepath: str) -> PolicyDocument:
        raise NotImplementedError


# ------------------------------------------------------------------
# HTML
# ------------------------------------------------------------------
class HTMLAdapter(BaseAdapter):
    def parse(self, filepath: str) -> PolicyDocument:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            raw = f.read()

        soup = BeautifulSoup(raw, 'html.parser')

        # Remove definitively-noise tags
        for tag in soup(NOISE_TAGS):
            tag.decompose()

        # Remove containers identified as noise by CSS class/id
        for tag in soup.find_all(True):
            cls = ' '.join(tag.get('class', []))
            id_ = tag.get('id', '')
            if NOISE_CLASS_RE.search(cls) or NOISE_CLASS_RE.search(id_):
                tag.decompose()

        # Remove link-dense sections (navigation / footer lists)
        # Heuristic: if >30% of tokens are anchor texts, it's a nav block
        for tag in soup.find_all(['ul', 'ol', 'div', 'section']):
            links = tag.find_all('a')
            all_words = re.findall(r'\b\w+\b', tag.get_text())
            if links and all_words and len(links) / len(all_words) > 0.25:
                tag.decompose()

        # Best-effort title extraction
        title = ''
        if soup.title:
            title = (soup.title.string or '').strip()
        if not title:
            h1 = soup.find('h1')
            title = h1.get_text(strip=True) if h1 else os.path.basename(filepath)

        text = soup.get_text(separator='\n\n')

        return PolicyDocument(
            policy_id=str(uuid.uuid4()), source_type='HTML',
            source_path=filepath, source_url='', title=title,
            raw_text=raw, cleaned_text=text,
            metadata=PolicyMetadata(extraction_date=datetime.datetime.now().isoformat()),
        )


# ------------------------------------------------------------------
# XML  (multi-schema)
# ------------------------------------------------------------------
class XMLAdapter(BaseAdapter):
    def parse(self, filepath: str) -> PolicyDocument:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            raw = f.read()

        soup = BeautifulSoup(raw, 'xml')

        policy_tag = soup.find('POLICY') or soup.find('policy')
        url = ''
        if policy_tag and policy_tag.has_attr('website_url'):
            url = policy_tag['website_url']

        text = self._extract_text(soup)

        return PolicyDocument(
            policy_id=str(uuid.uuid4()), source_type='XML',
            source_path=filepath, source_url=url,
            title=os.path.basename(filepath),
            raw_text=raw, cleaned_text=text,
            metadata=PolicyMetadata(extraction_date=datetime.datetime.now().isoformat()),
        )

    def _extract_text(self, soup) -> str:
        # Schema 1: ACL / OPP corpus — POLICY > SECTION > SUBTITLE + SUBTEXT
        sections = soup.find_all(re.compile(r'^section$', re.I))
        if sections:
            parts = []
            for sec in sections:
                subtitle = sec.find(re.compile(r'^subtitle$', re.I))
                subtext  = sec.find(re.compile(r'^subtext$',  re.I))
                if subtitle:
                    parts.append(subtitle.get_text(strip=True))
                if subtext:
                    parts.append(subtext.get_text(separator='\n\n').strip())
            if parts:
                return '\n\n'.join(parts)

        # Schema 2: SUBTEXT-only
        subtexts = soup.find_all(re.compile(r'^subtext$', re.I))
        if subtexts:
            return '\n\n'.join(
                s.get_text(separator=' ').strip() for s in subtexts if s.get_text(strip=True)
            )

        # Schema 3: Generic XML — p / para / paragraph / text / content tags
        paras = soup.find_all(re.compile(r'^(p|para|paragraph|text|content|body)$', re.I))
        if paras:
            return '\n\n'.join(
                p.get_text(separator=' ').strip() for p in paras if p.get_text(strip=True)
            )

        # Schema 4: Flat fallback
        return soup.get_text(separator='\n\n')


# ------------------------------------------------------------------
# PDF
# ------------------------------------------------------------------
class PDFAdapter(BaseAdapter):
    def parse(self, filepath: str) -> PolicyDocument:
        try:
            import pdfplumber
        except ImportError:
            raise RuntimeError('pdfplumber not installed. Run: pip install pdfplumber')

        pages_text = []
        first_lines: list = []
        last_lines: list = []

        # First pass: collect first / last lines per page to detect repeated headers/footers
        with pdfplumber.open(filepath) as pdf:
            n_pages = len(pdf.pages)
            for page in pdf.pages:
                raw = page.extract_text(x_tolerance=2, y_tolerance=2) or ''
                lines = [l.strip() for l in raw.splitlines() if l.strip()]
                if lines:
                    first_lines.append(lines[0])
                    last_lines.append(lines[-1])

            # Lines appearing on > 60% of pages → header/footer noise
            top_counts = Counter(first_lines)
            bot_counts = Counter(last_lines)
            repeat_noise: set = set()
            for line, cnt in {**top_counts, **bot_counts}.items():
                if n_pages > 0 and cnt / n_pages > 0.6:
                    repeat_noise.add(line)

            # Second pass: extract text, stripping repeated noise lines
            for page in pdf.pages:
                raw = page.extract_text(x_tolerance=2, y_tolerance=2) or ''
                filtered = [
                    l for l in raw.splitlines()
                    if l.strip() and l.strip() not in repeat_noise
                ]
                page_text = '\n'.join(filtered).strip()
                if page_text:
                    pages_text.append(page_text)

        full_text = '\n\n'.join(pages_text)
        title = os.path.basename(filepath).rsplit('.', 1)[0]

        return PolicyDocument(
            policy_id=str(uuid.uuid4()), source_type='PDF',
            source_path=filepath, source_url='', title=title,
            raw_text=full_text, cleaned_text=full_text,
            metadata=PolicyMetadata(extraction_date=datetime.datetime.now().isoformat()),
        )


# ------------------------------------------------------------------
# DOCX
# ------------------------------------------------------------------
class DOCXAdapter(BaseAdapter):
    def parse(self, filepath: str) -> PolicyDocument:
        try:
            from docx import Document
        except ImportError:
            raise RuntimeError('python-docx not installed. Run: pip install python-docx')

        doc = Document(filepath)
        parts = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                parts.append('')
                continue
            style_name = (para.style.name or '').lower()
            # Heading styles become section boundaries → wrap with blank lines
            if 'heading' in style_name:
                parts.append('')
                parts.append(text)
                parts.append('')
            else:
                parts.append(text)

        full_text = '\n'.join(parts)

        title = os.path.basename(filepath).rsplit('.', 1)[0]
        try:
            cp_title = doc.core_properties.title
            if cp_title:
                title = cp_title
        except Exception:
            pass

        return PolicyDocument(
            policy_id=str(uuid.uuid4()), source_type='DOCX',
            source_path=filepath, source_url='', title=title,
            raw_text=full_text, cleaned_text=full_text,
            metadata=PolicyMetadata(extraction_date=datetime.datetime.now().isoformat()),
        )


# ------------------------------------------------------------------
# TXT
# ------------------------------------------------------------------
class TXTAdapter(BaseAdapter):
    def parse(self, filepath: str) -> PolicyDocument:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            raw = f.read()
        return PolicyDocument(
            policy_id=str(uuid.uuid4()), source_type='TXT',
            source_path=filepath, source_url='',
            title=os.path.basename(filepath),
            raw_text=raw, cleaned_text=raw,
            metadata=PolicyMetadata(extraction_date=datetime.datetime.now().isoformat()),
        )


# =============================================================================
# 3. SEGMENTER
# =============================================================================

class PolicySegmenter:
    MIN_WORDS   = 5     # minimum words for a clause to be kept
    MAX_WORDS   = 50    # maximum words before a clause is forcibly re-split
    SENT_WINDOW = 1     # sentences grouped per clause (blob path)
    # Threshold: if fewer than 3 double-newlines, treat as blob
    BLOB_THRESHOLD = 3

    SECTION_KEYWORDS = {
        'Information Collection':      ['collect', 'information we collect', 'what we collect', 'gathering'],
        'Information Use':             ['how we use', 'use your information', 'purpose', 'utilize'],
        'Information Sharing':         ['share', 'sharing', 'disclosure', 'third part', 'disclose'],
        'Cookies & Tracking':          ['cookie', 'tracking', 'pixel', 'beacon', 'web beacon', 'fingerprint'],
        'Security':                    ['security', 'protect', 'safeguard', 'encryption', 'breach'],
        'Data Retention':              ['retain', 'retention', 'how long', 'storage period', 'deletion'],
        'User Rights':                 ['rights', 'access', 'delete', 'opt-out', 'choice', 'withdraw', 'erasure', 'portability'],
        "Children's Privacy":          ['children', 'kids', 'under 13', 'coppa', 'minor'],
        'International Transfers':     ['international', 'transfer', 'cross-border', 'adequacy', 'standard contractual'],
        'Contact Information':         ['contact', 'questions', 'reach us', 'inquiries', 'complaints'],
        'Policy Changes':              ['change', 'update', 'modify', 'amendment', 'revision', 'notify'],
        'Do Not Track':                ['do not track', 'dnt', 'opt out of tracking'],
        'Legal Basis':                 ['legal basis', 'legitimate interest', 'consent', 'legal ground'],
    }

    TAG_RULES = {
        'COLLECTION':  r'\b(collect|gather|obtain|receive|record)\b',
        'USE':         r'\b(use|utilize|process|purpose|employ)\b',
        'SHARING':     r'\b(share|disclose|transfer|third.part|vendor|partner)\b',
        'SECURITY':    r'\b(secure|protect|encryption|safeguard|breach|unauthorized)\b',
        'COOKIES':     r'\b(cookie|pixel|tracking|beacon|fingerprint|identifier)\b',
        'CHILDREN':    r'\b(children|kids|under\s*13|coppa|minor)\b',
        'RETENTION':   r'\b(retain|store|keep|delete|purge|period|duration)\b',
        'RIGHTS':      r'\b(right|access|erasure|portability|object|withdraw|opt.out)\b',
        'LEGAL_BASIS': r'\b(consent|legitimate\s+interest|legal\s+basis|contractual)\b',
    }

    def detect_section_name(self, text: str) -> str:
        text_lower = text.lower()
        for norm_name, kw_list in self.SECTION_KEYWORDS.items():
            if any(kw in text_lower for kw in kw_list):
                return norm_name
        return 'General Information'

    def assign_tags(self, text: str) -> list:
        text_lower = text.lower()
        return [
            tag for tag, pattern in self.TAG_RULES.items()
            if re.search(pattern, text_lower)
        ]

    def _is_heading(self, block: str) -> bool:
        words = block.split()
        return (
            1 <= len(words) <= 10
            and len(block) < 80
            and not block.endswith('.')
            and (
                block.istitle()
                or block.isupper()
                or bool(re.match(r'^\d+[\.\)]\s+[A-Z]', block))
                or (words[0][0].isupper() and len(words) <= 5)
            )
        )

    def _build_clause(self, text: str, policy_id: str, counter: int,
                      section_name: str, offset: int) -> Clause:
        return Clause(
            clause_id=f'{policy_id}-C{counter}',
            section_name=section_name,
            clause_text=text,
            start_offset=offset,
            end_offset=offset + len(text),
            preprocessing_tags=self.assign_tags(text),
        )

    def segment(self, doc: PolicyDocument) -> PolicyDocument:
        text = doc.cleaned_text
        seen_fingerprints: set = set()
        clause_counter = 1
        offset = 0

        current_section = Section(
            original_heading='General',
            normalized_heading='General Information',
            section_text='', start_offset=0, end_offset=0,
        )

        def _fingerprint(s: str) -> str:
            return re.sub(r'\s+', ' ', s[:80].lower().strip())

        def _try_add(candidate: str):
            nonlocal clause_counter, offset
            candidate = candidate.strip()
            if len(candidate.split()) < self.MIN_WORDS:
                return
            fp = _fingerprint(candidate)
            if fp in seen_fingerprints:
                return
            seen_fingerprints.add(fp)

            # If still too long, recursively re-split on fullstops
            if len(candidate.split()) > self.MAX_WORDS:
                sents = _split_sentences(candidate)
                for sub in _group_sentences(sents, window=1, max_words=self.MAX_WORDS):
                    sub = sub.strip()
                    if len(sub.split()) < self.MIN_WORDS:
                        continue
                    fp2 = _fingerprint(sub)
                    if fp2 in seen_fingerprints:
                        continue
                    seen_fingerprints.add(fp2)
                    c = self._build_clause(sub, doc.policy_id, clause_counter,
                                           current_section.normalized_heading, offset)
                    current_section.clauses.append(c)
                    doc.clauses.append(c)
                    clause_counter += 1
                    offset += len(sub) + 2
                return

            c = self._build_clause(candidate, doc.policy_id, clause_counter,
                                   current_section.normalized_heading, offset)
            current_section.clauses.append(c)
            doc.clauses.append(c)
            clause_counter += 1
            offset += len(candidate) + 2

        # ── Determine path: structured vs. blob ───────────────────────────────
        double_newline_count = text.count('\n\n')
        is_blob = double_newline_count < self.BLOB_THRESHOLD

        if is_blob:
            # BLOB PATH ────────────────────────────────────────────────────────
            # The document has no paragraph structure.
            # Split into sentences on fullstops, then group into clause chunks.
            sentences = _split_sentences(text)
            for candidate in _group_sentences(sentences, self.SENT_WINDOW, self.MAX_WORDS):
                _try_add(candidate)
        else:
            # STRUCTURED PATH ─────────────────────────────────────────────────
            # Paragraph-level split, with subclause refinement.
            blocks = re.split(r'\n\n+', text)

            for block in blocks:
                block = block.strip()
                if not block:
                    continue

                # Section heading detection
                if self._is_heading(block):
                    if current_section.clauses:
                        doc.sections.append(current_section)
                    norm_heading = self.detect_section_name(block)
                    current_section = Section(
                        original_heading=block,
                        normalized_heading=norm_heading,
                        section_text='',
                        start_offset=offset,
                        end_offset=offset + len(block),
                    )
                    offset += len(block) + 2
                    continue

                # Sub-clause split: numbered items, bullets, semicolon lists
                sub_clauses = re.split(
                    r'(?:\n\s*[-*•]\s+)'           # bullet point
                    r'|(?:\n\s*\d+[\.\)]\s+)'      # numbered list
                    r'|(?:;\s+(?:and\s+|or\s+)?(?=[A-Z]))',  # semicolon before capitalised clause
                    block,
                )

                for sc in sub_clauses:
                    sc = sc.strip()
                    if not sc:
                        continue

                    # Long sub-clause → fullstop split within it
                    if len(sc.split()) > self.MAX_WORDS:
                        sents = _split_sentences(sc)
                        for candidate in _group_sentences(sents, self.SENT_WINDOW, self.MAX_WORDS):
                            _try_add(candidate)
                    else:
                        _try_add(sc)

        if current_section.clauses:
            doc.sections.append(current_section)

        return doc


# =============================================================================
# 4. NLP ENRICHER
# =============================================================================

class NLPEnricher:
    def __init__(self):
        try:
            spacy.require_gpu()
        except Exception:
            pass
        self.nlp = spacy.load('en_core_web_sm')

    def enrich(self, doc: PolicyDocument) -> PolicyDocument:
        if not doc.cleaned_text:
            return doc
        spacy_doc = self.nlp(doc.cleaned_text[:100_000])
        for ent in spacy_doc.ents:
            if ent.label_ in ('ORG', 'GPE', 'LOC', 'DATE'):
                doc.metadata.extracted_entities.append(Entity(
                    text=ent.text, label=ent.label_,
                    start_offset=ent.start_char, end_offset=ent.end_char,
                ))
        return doc


# =============================================================================
# 5. QUALITY VALIDATOR
# =============================================================================

class QualityValidator:
    @staticmethod
    def validate(doc: PolicyDocument) -> PolicyDocument:
        word_count = len(doc.cleaned_text.split())
        doc.metadata.word_count = word_count

        score = 100
        flags = []

        if word_count == 0:
            flags.append('EMPTY_DOCUMENT')
            score = 0
        elif word_count < 200:
            flags.append('INSUFFICIENT_TEXT')
            score -= 50

        if not doc.clauses:
            flags.append('SEGMENTATION_FAILURE')
            score -= 30
        elif word_count > 500 and len(doc.clauses) < 3:
            # Large doc with very few clauses — likely blob segmentation still failed
            flags.append('POSSIBLE_BLOB_FAILURE')
            score -= 20

        if not doc.sections:
            flags.append('SECTION_DETECTION_FAILURE')
            score -= 10

        doc.validation_flags = flags
        doc.quality_score = max(0, score)
        return doc


# =============================================================================
# 6. PIPELINE ORCHESTRATOR
# =============================================================================

class PolicyIngestionPipeline:
    def __init__(self):
        self.adapters = {
            '.html':  HTMLAdapter(),
            '.htm':   HTMLAdapter(),
            '.xml':   XMLAdapter(),
            '.txt':   TXTAdapter(),
            '.pdf':   PDFAdapter(),
            '.docx':  DOCXAdapter(),
        }
        self.segmenter    = PolicySegmenter()
        self.nlp_enricher = NLPEnricher()
        self.validator    = QualityValidator()

    def process_file(self, filepath: str) -> PolicyDocument:
        ext = os.path.splitext(filepath)[1].lower()
        adapter = self.adapters.get(ext, TXTAdapter())

        # 1. Format-specific extraction
        doc = adapter.parse(filepath)

        # 2. Universal noise cleaning
        doc.cleaned_text = TextCleaner.clean(doc.cleaned_text)

        # 3. Clause segmentation
        doc = self.segmenter.segment(doc)

        # 4. NLP entity enrichment
        doc = self.nlp_enricher.enrich(doc)

        # 5. Quality validation
        doc = self.validator.validate(doc)

        return doc
